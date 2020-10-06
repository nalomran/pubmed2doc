#!/usr/bin/env python3

"""Pubmed To Document

Created on Sep, 18 2020

@author: Nawaf Alomran

This module allows the user to write PubMed search results to
either Word or PDF with two display options i.e citation or listview.

Please visit the README file for more information about the module and some of
its advantages.

The user needs to supply a query to search against the PubMed database and an
email address to access the PubMed database. The list of arguments for the
the module are shown below:

Input + Options
----------------

    + -q: a query to be searched against PubMed database (REQUIRED)

    + -e: a user's email to access to the PubMed db (REQUIRED)

    + -pdf: write PubMed results to PDF (OPTIONAL) (Default value is True)

    + -word: write PubMed results to Word (OPTIONAL) (Default value is False)

    + -retmax: total num of records from query to be retrieved (OPTIONAL)
    (Default value is 20)

    + -sopt: type of returned results display options (citation or listview)
    (OPTIONAL) (Default is citation)

    + -mndate: custom start or minimum publication date (OPTIONAL)"

    + -mxdate: custom end or maximum publication date (OPTIONAL)"

    + is_abstract: include abstract to your search results



Output
------
    + a Word or PDF document consists of the results from PubMed according to
    the type of display option chosen by the user


How to Run
-----------

    # to write to a Word file with the display option listview for the
    query "gene expression"

    python pubmed2doc.py \
    -q "gene expression" \
    -e "your_email" \
    -pdf F \
    -word T \
    -sopt listview

    # to write to a Word file with the display option citation for the query
    "gene expression"

    python pubmed2doc.py \
    -q "gene expression" \
    -e "your_email" \
    -pdf F \
    -word T \
    -sopt citation
"""

import argparse
import os
import sys
from pathlib import Path

import docx  # type: ignore
import fpdf
from Bio import Entrez, Medline  # type: ignore
from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore
from fpdf import FPDF  # type: ignore

fpdf.set_global("SYSTEM_TTFONTS",
                os.path.join(os.path.dirname(__file__),
                             'font'))


def query_search_pubmed(query: str,
                        ret_max: str,
                        email: str,
                        min_date: str,
                        max_date: str
                        ):
    """Search PubMed via the user's query supplied through the command line

    Parameters
    ----------
    query: a query to be searched against PubMed database

    email: a user's email to access to the PubMed database

    ret_max: total number of records from query to be retrieved

    min_date: the minimum or start date to search

    max_date: the maximum or end date to search


    Return
    -------
    retrieve document summaries as records

    """

    Entrez.email = email

    if min_date and max_date:
        # search the PubMed db for the entered query
        search = Entrez.esearch(
            db="pubmed",
            term=query,
            sort="relevance",
            retmode="text",
            retmax=ret_max,
            mindate=min_date,
            maxdate=max_date
        )
    else:
        # search the PubMed db for the entered query
        search = Entrez.esearch(
            db="pubmed",
            term=query,
            sort="relevance",
            retmode="text",
            retmax=ret_max,
            usehistory='y'
        )

    search_records = Entrez.read(search)
    search.close()

    # get the list of ids for the searched records
    list_ids = search_records['IdList']

    print(f"\nTotal of {len(list_ids)} records retrieved!")

    ids = ",".join(list_ids)

    # return document summaries as a result handle
    fetch_records = Entrez.efetch(
        db="pubmed",
        id=ids,
        rettype="Medline",
        retmode="text",
        webenv=search_records['WebEnv'],
        query_key=search_records['QueryKey']

    )

    search_results = Medline.parse(fetch_records)
    # fetch_records.close()

    return search_results


def records_iterator(pubmed_results, is_abstract: bool):
    """Iterates over the records returned from PubMed results

    Parameters
    ----------
    pubmed_results: PubMed results stored in Bio.Entrez.Parser.ListElement

    is_abstract: include abstract to your search results

    Return
    -------
    generator of the PubMed results

    """

    for paper in pubmed_results:
        # handle Unicode Encode Error due to some unusual characters in the
        # authors names for PDF only
        authors = ", ".join(list(paper['AU']))

        title = paper['TI']

        abstract = paper['AB']

        journal = paper['JT']

        pub_date = paper['DP']

        so = paper['SO'].split('doi')

        vol_issue = f"{so[0].split()[-1]}"

        doi = so[-1]

        pmid = f"PMID: {paper['PMID']}"

        # handle pmcid in case if it is unavailable
        try:
            pmcid = f"pmcid: {paper['PMC']}"

        except KeyError:
            pmcid = 'pmcid: Not Available'

        if is_abstract:
            yield (authors, title, journal, pub_date, vol_issue, doi,
                   pmid, pmcid, abstract)

        else:
            yield (authors, title, journal, pub_date, vol_issue, doi,
                   pmid, pmcid)


def write_to_pdf(
        pubmed_results,
        style_method: str,
        query: str,
        is_abstract: bool) -> None:
    """Write the PubMed results to PDF

    Parameters
    ----------
    pubmed_results: PubMed results stored in Bio.Entrez.Parser.ListElement

    style_method: the style to be written to a PDF file
    (citation or listview)

    query: a query to be searched against PubMed database

    is_abstract: include abstract to your search results

    Return
    -------
    None
    """

    pdf_doc = FPDF()

    pdf_doc.add_page()

    # add font
    pdf_doc.add_font("NotoSans", style="",
                     fname="NotoSans-Regular.ttf",
                     uni=True)

    # setting the font name and size
    pdf_doc.set_font("NotoSans", size=12)

    # configure the header of the PDF document
    pdf_doc.cell(200, 20, txt=f"Search Results for {query.title()}",
                 ln=1, align='C')

    pdf_doc.set_display_mode(zoom='real')

    records = records_iterator(pubmed_results, is_abstract)

    for i, (authors, title, journal, *pub_info) in enumerate(records, 1):
        try:
            pub_date, vol_issue, doi, pmid, pmcid, abstract = pub_info
        except ValueError:
            pub_date, vol_issue, doi, pmid, pmcid = pub_info

        url = f"http://www.ncbi.nlm.nih.gov/pubmed/" + pmid.split()[-1]

        if style_method == 'citation':
            # the following add the information retrieved from the PubMed
            # results
            pdf_doc.set_text_color(0, 0, 0)
            pdf_doc.multi_cell(180, 7, txt=str(i) + ': ' +
                                           authors + '. ' +
                                           title + ' ' +
                                           journal + '. ' +
                                           pub_date + ';' +
                                           vol_issue +
                                           " doi" + doi + ' ' +
                                           pmid + '; ' +
                                           pmcid + '. ' +
                                           url + '.'
                               )
            pdf_doc.set_text_color(0, 0, 255)

            # add line break
            pdf_doc.ln(h='2')

        # selecting the other style "the listview"
        else:
            pdf_doc.set_text_color(0, 0, 0)
            pdf_doc.multi_cell(180, 7, txt=str(i) + ': ' + "Authors: " +
                                           authors)

            pdf_doc.multi_cell(180, 7, txt="Title: " + title)

            pdf_doc.cell(90, 7, txt="Journal Name: " + journal,
                         ln=1,
                         align='L')

            pdf_doc.cell(50, 7, txt="Publication Date: " +
                                    pub_date,
                         ln=2,
                         align='L')

            pdf_doc.cell(50, 7, txt="Volume, Issue, pages: " + vol_issue,
                         ln=2,
                         align='L')

            pdf_doc.cell(50, 7, txt="doi" + doi, ln=1, align='L')

            pdf_doc.cell(50, 7, txt=pmid, ln=2, align='L')

            pdf_doc.cell(50, 7, txt=pmcid, ln=2, align='L')

            pdf_doc.cell(50, 7, txt="url: " + url, ln=9, align='L')

            try:
                pdf_doc.multi_cell(190, 7, txt="Abstract: " + abstract,
                                   align='L')
            except UnboundLocalError:
                pdf_doc.ln(h='2')
                continue

            pdf_doc.ln(h='2')

    # write the information above to PDF file
    print("\nwriting to a PDF file...")
    pdf_doc.output('output/PubMed_Results.pdf')
    print("\nDone writing.")


def add_hyperlink(paragraph, url, text, color, underline):
    """ A function that places a hyperlink within a paragraph object.
    This function is taking from the github issue in the following link:
    # https://github.com/python-openxml/python-docx/issues/74 for pyhton-docx

    Parameters
    ----------
    paragraph: The paragraph we are adding the hyperlink to.
    url: A string containing the required url
    text: The text displayed for the url

    Return
    -------
    The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation
    # id value
    part = paragraph.part
    r_id = part.relate_to(url,
                          docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                          is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add the required text to the
    # w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def write_to_word(pubmed_results,
                  style_method: str,
                  query: str,
                  is_abstract: bool) -> None:
    """Write the PubMed results to Word

    Parameters
    ----------
    pubmed_results: PubMed results stored in Bio.Entrez.Parser.ListElement

    style_method: the style to be written to a PDF file
    (citation or listview)

    query: a query to be searched against PubMed database

    is_abstract: include abstract to your search results

    Return
    -------
    None
    """

    document = Document()

    # styling document (font name and size)
    style = document.styles['Normal']

    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # handling header configurations
    header = document.add_paragraph()
    header.alignment = 1  # center the title
    header.add_run(
        f'PubMed Search Results for {query.title()}')  # .bold = True

    records = records_iterator(pubmed_results, is_abstract)

    for i, (authors, title, journal, *pub_info) in enumerate(records, 1):

        try:
            pub_date, vol_issue, doi, pmid, pmcid, abstract = pub_info
        except ValueError:
            pub_date, vol_issue, doi, pmid, pmcid = pub_info

        url = f"http://www.ncbi.nlm.nih.gov/pubmed/" + pmid.split()[-1]

        if style_method == 'citation':

            paragraph = document.add_paragraph(str(i) + ': ' +
                                               authors + '. ' +
                                               title + ' ' +
                                               journal + '. ' +
                                               pub_date + ';' +
                                               vol_issue + ' ' +
                                               "doi" + doi + ' ' +
                                               pmid + '; ' +
                                               pmcid + '.'
                                               )

            # to include the url as part of the citation
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(1)

            # for url hyperlink
            add_hyperlink(document.add_paragraph(), url, url, 'blue',
                          True)

        else:

            document.add_paragraph(str(i) + ":" + "  Authors: " + authors)

            document.add_paragraph("Title: " + title)

            document.add_paragraph(
                "Journal Name: " + journal)

            document.add_paragraph("Publication Date: " + pub_date)

            document.add_paragraph("Volume, Issue, Pages: " + vol_issue)

            document.add_paragraph("doi" + doi)

            document.add_paragraph(pmid)

            document.add_paragraph(pmcid)

            add_hyperlink(document.add_paragraph('Url: '),
                          url, url, 'blue', True)

            try:
                document.add_paragraph("Abstract: " + abstract)
            except UnboundLocalError:
                continue

    # write the information above to MS Word file
    print("\nwriting to a MS Word file...")
    document.save('output/PubMed_Results.docx')
    print("\nDone writing.")


def bool_conv_args(args: str) -> bool:
    """Convert string argument value to boolean True or False

    Parameters
    ----------
    args: argument value to represent True or False


    Return
    -------
    a converted string argument to a boolean value
    """

    # consider most possible truth values scenarios supplied by the user
    if args.lower() in ['yes', 'true', 't', 'y']:
        return True

    elif args.lower() in ['no', 'false', 'f', 'n']:
        return False

    else:
        raise argparse.ArgumentParser('Please make sure to enter a boolean '
                                      'value i.e. True or False.')


def main(args) -> None:
    """Direct the operations and processes in this module

    Parameters
    ----------
    args: argparse Namespace class that consists of argument-value pairs


    Return
    -------
    None
    """

    query = args.query

    ret_max = args.ret_max

    email = args.email

    style_method = args.style_opt

    min_date_arg = args.min_date

    max_date_arg = args.max_date

    results = query_search_pubmed(query, ret_max, email,
                                  min_date_arg, max_date_arg)

    # create a directory named "output" if it doesn't exists
    if not Path("output").exists():
        print(f"\ncreating output directory named 'output'...")
        Path('output').mkdir(parents=True, exist_ok=True)

    to_pdf = args.to_pdf

    to_word = args.to_word

    is_abstract = args.is_abstract

    if to_word:
        write_to_word(results, style_method, query, is_abstract)
        to_pdf = False  # PDF writing is disabled

    if to_pdf:
        write_to_pdf(results, style_method, query, is_abstract)


def run_command_lines() -> None:
    """Add and parse the arguments from the command line and execute the main
    function

    Parameters
    ----------
    None


    Return
    -------
    None
    """

    USAGE = '''write PubMed results to Word or PDF file'''

    parser = argparse.ArgumentParser(description=USAGE)
    parser.add_argument('-q',
                        dest="query",
                        required=True,
                        help="a query to be searched against PubMed database")

    parser.add_argument('-e',
                        dest="email",
                        required=True,
                        help="a user's email to access to the PubMed db")

    parser.add_argument('-pdf',
                        dest="to_pdf",
                        type=bool_conv_args,
                        default=True,
                        help="write Pubmed results to PDF (OPTIONAL) "
                             "(Default value is True)")

    parser.add_argument('-word',
                        dest="to_word",
                        type=bool_conv_args,
                        default=False,
                        help="write Pubmed results to Word (OPTIONAL) "
                             "(Default value is False)")

    parser.add_argument('-retmax',
                        dest="ret_max",
                        default=20,
                        help="total num of records from query to be retrieved "
                             "(OPTIONAL) (Default is 20)")

    parser.add_argument('-sopt',
                        dest="style_opt",
                        default="citation",
                        help="type of returned results style format "
                             "(citation or listview) (OPTIONAL) "
                             "(Default is citation)")

    parser.add_argument('-mndate',
                        dest="min_date",
                        default=None,
                        help="custom start or minimum publication date "
                             "(OPTIONAL)")

    parser.add_argument('-mxdate',
                        dest="max_date",
                        default=None,
                        help="custom end or maximum publication date "
                             "(OPTIONAL)")

    parser.add_argument('-abs',
                        dest="is_abstract",
                        type=bool_conv_args,
                        default=False,
                        help="include abstract to your search results"
                             "(OPTIONAL)")

    args = parser.parse_args()

    # execute the main functions with the command line arguments
    try:
        main(args)
    except KeyboardInterrupt:
        sys.exit('\nprogram is terminated by the user!')


if __name__ == '__main__':
    run_command_lines()
